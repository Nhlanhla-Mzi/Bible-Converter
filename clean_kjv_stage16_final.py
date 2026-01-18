import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ================= FILES =================
INPUT_FILE = "kjv_formatted.txt"
OUTPUT_FILE = "KJV_Cleaned_Final.docx"

doc = Document()

# ================= BASE STYLE =================
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# 🔒 FORCE HEADING 1 STYLE (TOC SAFE)
heading1 = doc.styles['Heading 1']
heading1.font.name = 'Times New Roman'
heading1.font.color.rgb = None  # force black

# ================= FONT SIZES =================
CHAPTER_FONT_SIZE = Pt(20)
BOOK_DESCRIPTOR_SIZE = Pt(14)
BOOK_NAME_SIZE = Pt(26)

# ================= BOOK TITLES =================
BOOK_TITLES = {
    "Genesis": "The First Book of Moses, Called",
    "Exodus": "The Second Book of Moses, Called",
    "Leviticus": "The Third Book of Moses, Called",
    "Numbers": "The Fourth Book of Moses, Called",
    "Deuteronomy": "The Fifth Book of Moses, Called",
    "Joshua": "The Book of",
    "Judges": "The Book of",
    "Ruth": "The Book of",
    "1 Samuel": "The First Book of",
    "2 Samuel": "The Second Book of",
    "1 Kings": "The First Book of the Kings",
    "2 Kings": "The Second Book of the Kings",
    "1 Chronicles": "The First Book of the Chronicles",
    "2 Chronicles": "The Second Book of the Chronicles",
    "Ezra": "The Book of",
    "Nehemiah": "The Book of",
    "Esther": "The Book of",
    "Job": "The Book of",
    "Psalms": "The Book of",
    "Proverbs": "The Proverbs",
    "Ecclesiastes": "",
    "Song of Solomon": "The Song of",
    "Isaiah": "The Book of the Prophet",
    "Jeremiah": "The Book of the Prophet",
    "Lamentations": "The Lamentations of",
    "Ezekiel": "The Book of the Prophet",
    "Daniel": "The Book of",
    "Hosea": "The Book of",
    "Joel": "The Book of",
    "Amos": "The Book of",
    "Obadiah": "The Book of",
    "Jonah": "The Book of",
    "Micah": "The Book of",
    "Nahum": "The Book of",
    "Habakkuk": "The Book of",
    "Zephaniah": "The Book of",
    "Haggai": "The Book of",
    "Zechariah": "The Book of",
    "Malachi": "The Book of",
    "Matthew": "The Gospel According to Saint",
    "Mark": "The Gospel According to Saint",
    "Luke": "The Gospel According to Saint",
    "John": "The Gospel According to Saint",
    "Acts": "The Acts of the Apostles",
    "Romans": "The Epistle of Paul the Apostle to the",
    "1 Corinthians": "The First Epistle of Paul the Apostle to the",
    "2 Corinthians": "The Second Epistle of Paul the Apostle to the",
    "Galatians": "The Epistle of Paul the Apostle to the",
    "Ephesians": "The Epistle of Paul the Apostle to the",
    "Philippians": "The Epistle of Paul the Apostle to the",
    "Colossians": "The Epistle of Paul the Apostle to the",
    "1 Thessalonians": "The First Epistle of Paul the Apostle to the",
    "2 Thessalonians": "The Second Epistle of Paul the Apostle to the",
    "1 Timothy": "The First Epistle of Paul the Apostle to",
    "2 Timothy": "The Second Epistle of Paul the Apostle to",
    "Titus": "The Epistle of Paul the Apostle to",
    "Philemon": "The Epistle of Paul the Apostle to",
    "Hebrews": "The Epistle of Paul the Apostle to the",
    "James": "The General Epistle of",
    "1 Peter": "The First Epistle General of",
    "2 Peter": "The Second Epistle General of",
    "1 John": "The First Epistle General of",
    "2 John": "The Second Epistle General of",
    "3 John": "The Third Epistle General of",
    "Jude": "The General Epistle of",
    "Revelation": "The Revelation of Saint John the Divine"
}

# ================= REGEX =================
JUNK_START = re.compile(r'^[\u25A0\u25A1\uFFFD\s]+')
REMOVE_KJV_ONLINE = re.compile(r'KJV[\s_]*Online', re.IGNORECASE)
CHAPTER_ONLY = re.compile(r'^\d+$')
VERSE_LINE = re.compile(r'^(\d+)([\u202F\u00A0\s]+)(.*)')
UNDERSCORE_ONLY = re.compile(r'^_+$')
BOOK_TITLE_DASHED = re.compile(r'^-+\s*(.+?)\s*-+$')
BOOK_CHAPTER_LINE = re.compile(r'^(.+?)\s+(\d+)$')

# ================= HELPERS =================
def add_toc(document):
    p = document.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(ns.qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-1" \\h \\z \\u'
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def add_text_with_italics(p, text):
    parts = re.split(r'(_[^_]+_)', text)
    for part in parts:
        if part.startswith('_') and part.endswith('_'):
            r = p.add_run(part[1:-1])
            r.italic = True
        else:
            p.add_run(part)

# ================= COVER PAGE =================
cover = doc.add_paragraph("THE HOLY BIBLE\n\nKing James Version")
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.runs[0].bold = True
cover.runs[0].font.size = Pt(30)
doc.add_page_break()

# ================= TOC PAGE =================
toc_title = doc.add_paragraph("CONTENTS")
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_title.runs[0].bold = True
toc_title.runs[0].font.size = Pt(20)

add_toc(doc)
doc.add_page_break()

# ================= INTRODUCTION =================
intro = doc.add_paragraph(
    "INTRODUCTION\n\n"
    "This edition of the Holy Bible presents the text of the King James Version "
    "in a clean, readable, and structured format. Book titles follow the traditional "
    "KJV naming conventions. Chapter numbers are visually distinguished, and "
    "section headings are clearly marked.\n\n"
    "This document preserves the original language, structure, and emphasis of the "
    "Authorized Version, while applying modern typographic clarity for extended reading "
    "and study."
)
intro.runs[0].bold = True
doc.add_page_break()

# ================= MAIN PROCESS =================
pending_chapter = None

with open(INPUT_FILE, 'r', encoding='utf-8') as f:
    for raw in f:
        line = raw.rstrip()
        line = REMOVE_KJV_ONLINE.sub('', line)
        line = JUNK_START.sub('', line)

        if not line.strip() or UNDERSCORE_ONLY.match(line):
            continue

        dashed = BOOK_TITLE_DASHED.match(line)
        if dashed:
            book = dashed.group(1).strip()
            if book in BOOK_TITLES:
                descriptor = BOOK_TITLES[book]

                if descriptor:
                    p1 = doc.add_paragraph(descriptor)
                    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p1.runs[0].font.size = BOOK_DESCRIPTOR_SIZE

                p2 = doc.add_heading(book.upper(), level=1)
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p2.runs[0].bold = True
                p2.runs[0].font.size = BOOK_NAME_SIZE
                continue

        if CHAPTER_ONLY.match(line):
            pending_chapter = line
            continue

        bc = BOOK_CHAPTER_LINE.match(line)
        if bc:
            pending_chapter = bc.group(2)
            continue

        verse = VERSE_LINE.match(line)
        if verse:
            vnum, space, text = verse.groups()
            p = doc.add_paragraph()
            if vnum == "1" and pending_chapter:
                r = p.add_run(pending_chapter)
                r.bold = True
                r.font.size = CHAPTER_FONT_SIZE
                add_text_with_italics(p, space + text)
                pending_chapter = None
            else:
                add_text_with_italics(p, line)
            continue

        p = doc.add_paragraph()
        add_text_with_italics(p, line)
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(12.5)

doc.save(OUTPUT_FILE)
print("Finished: KJV_Cleaned_Final.docx created")
