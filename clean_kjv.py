import re
from docx import Document
from docx.shared import Pt

INPUT_FILE = "kjv_formatted.txt"
OUTPUT_FILE = "KJV_Cleaned_Final.docx"

doc = Document()

# Optional: set default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# Unicode junk often seen before sentences
JUNK_CHARS = r'^[\u25A0\u25A1\uFFFD\s]+'  # ■ □ �

# Regex patterns
ITALICS_PATTERN = re.compile(r'_(.*?)_')
CHAPTER_PATTERN = re.compile(
    r'^(Genesis|Exodus|Leviticus|Numbers|Deuteronomy|Joshua|Judges|Ruth|'
    r'1 Samuel|2 Samuel|1 Kings|2 Kings|1 Chronicles|2 Chronicles|'
    r'Ezra|Nehemiah|Esther|Job|Psalms|Proverbs|Ecclesiastes|Song of Solomon|'
    r'Isaiah|Jeremiah|Lamentations|Ezekiel|Daniel|Hosea|Joel|Amos|Obadiah|'
    r'Jonah|Micah|Nahum|Habakkuk|Zephaniah|Haggai|Zechariah|Malachi|'
    r'Matthew|Mark|Luke|John|Acts|Romans|1 Corinthians|2 Corinthians|'
    r'Galatians|Ephesians|Philippians|Colossians|'
    r'1 Thessalonians|2 Thessalonians|1 Timothy|2 Timothy|Titus|Philemon|'
    r'Hebrews|James|1 Peter|2 Peter|1 John|2 John|3 John|Jude|Revelation)'
    r'\s+(\d+)$'
)

with open(INPUT_FILE, 'r', encoding='utf-8') as f:
    for line in f:
        line = line.rstrip()

        # Remove junk squares at sentence start
        line = re.sub(JUNK_CHARS, '', line)

        # Normalize chapter headers
        chap_match = CHAPTER_PATTERN.match(line)
        if chap_match:
            p = doc.add_paragraph()
            run = p.add_run(chap_match.group(2))
            run.bold = True
            p.paragraph_format.space_before = Pt(12)
            continue

        # Process italics
        p = doc.add_paragraph()
        last_idx = 0

        for m in ITALICS_PATTERN.finditer(line):
            p.add_run(line[last_idx:m.start()])
            ital = p.add_run(m.group(1))
            ital.italic = True
            last_idx = m.end()

        p.add_run(line[last_idx:])

doc.save(OUTPUT_FILE)

print("Finished: KJV_Cleaned.docx created")
