import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns

# ================= FILES =================
INPUT_FILE = "kjv_formatted.txt"
OUTPUT_FILE = "KJV_Cleaned_Final.docx"

# ================= DOCUMENT =================
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# ================= SIZES =================
CHAPTER_FONT_SIZE = Pt(20)
BOOK_NAME_SIZE = Pt(26)

# ================= BOOK TITLES =================
BOOK_TITLES = {
    "Genesis": "The First Book of Moses, Called",
    "Exodus": "The Second Book of Moses, Called",
    "Leviticus": "The Third Book of Moses, Called",
    "Numbers": "The Fourth Book of Moses, Called",
    "Deuteronomy": "The Fifth Book of Moses, Called",

    "Joshua": "The Book of",
    "Judges": "The Book of",
    "Ruth": "The Book of",
    "1 Samuel": "The First Book of",
    "2 Samuel": "The Second Book of",
    "1 Kings": "The First Book of the Kings",
    "2 Kings": "The Second Book of the Kings",
    "1 Chronicles": "The First Book of the Chronicles",
    "2 Chronicles": "The Second Book of the Chronicles",
    "Ezra": "The Book of",
    "Nehemiah": "The Book of",
    "Esther": "The Book of",

    "Job": "The Book of",
    "Psalms": "The Book of",
    "Proverbs": "The Proverbs",
    "Ecclesiastes": "",
    "Song of Solomon": "The Song of",

    "Isaiah": "The Book of the Prophet",
    "Jeremiah": "The Book of the Prophet",
    "Lamentations": "The Lamentations of",
    "Ezekiel": "The Book of the Prophet",
    "Daniel": "The Book of",

    "Hosea": "The Book of",
    "Joel": "The Book of",
    "Amos": "The Book of",
    "Obadiah": "The Book of",
    "Jonah": "The Book of",
    "Micah": "The Book of",
    "Nahum": "The Book of",
    "Habakkuk": "The Book of",
    "Zephaniah": "The Book of",
    "Haggai": "The Book of",
    "Zechariah": "The Book of",
    "Malachi": "The Book of",

    "Matthew": "The Gospel According to Saint",
    "Mark": "The Gospel According to Saint",
    "Luke": "The Gospel According to Saint",
    "John": "The Gospel According to Saint",
    "Acts": "The Acts of the Apostles",

    "Romans": "The Epistle of Paul the Apostle to the",
    "1 Corinthians": "The First Epistle of Paul the Apostle to the",
    "2 Corinthians": "The Second Epistle of Paul the Apostle to the",
    "Galatians": "The Epistle of Paul the Apostle to the",
    "Ephesians": "The Epistle of Paul the Apostle to the",
    "Philippians": "The Epistle of Paul the Apostle to the",
    "Colossians": "The Epistle of Paul the Apostle to the",
    "1 Thessalonians": "The First Epistle of Paul the Apostle to the",
    "2 Thessalonians": "The Second Epistle of Paul the Apostle to the",
    "1 Timothy": "The First Epistle of Paul the Apostle to",
    "2 Timothy": "The Second Epistle of Paul the Apostle to",
    "Titus": "The Epistle of Paul the Apostle to",
    "Philemon": "The Epistle of Paul the Apostle to",
    "Hebrews": "The Epistle of Paul the Apostle to the",

    "James": "The General Epistle of",
    "1 Peter": "The First Epistle General of",
    "2 Peter": "The Second Epistle General of",
    "1 John": "The First Epistle General of",
    "2 John": "The Second Epistle General of",
    "3 John": "The Third Epistle General of",
    "Jude": "The General Epistle of",

    "Revelation": "The Revelation of Saint John the Divine"
}

# ================= REGEX =================
REMOVE_KJV_ONLINE = re.compile(r'KJV[\s_]*Online', re.IGNORECASE)
JUNK_START = re.compile(r'^[\u25A0\u25A1\uFFFD\s]+')
CHAPTER_ONLY = re.compile(r'^\d+$')
VERSE_LINE = re.compile(r'^(\d+)([\u202F\u00A0\s]+)(.*)')
UNDERSCORE_ONLY = re.compile(r'^_+$')

pending_chapter = None

# ================= PROCESS =================
with open(INPUT_FILE, "r", encoding="utf-8") as f:
    for raw in f:
        line = raw.rstrip()

        line = REMOVE_KJV_ONLINE.sub("", line)
        line = JUNK_START.sub("", line)

        if not line.strip():
            continue

        if UNDERSCORE_ONLY.match(line):
            continue

        # -------- Book titles (ORDER IS PRESERVED) --------
        if line in BOOK_TITLES:
            # Descriptor line
            p1 = doc.add_paragraph()
            r1 = p1.add_run(BOOK_TITLES[line])
            r1.font.size = CHAPTER_FONT_SIZE
            r1.bold = False
            p1.paragraph_format.alignment = 1
            p1.paragraph_format.space_before = Pt(36)
            p1.paragraph_format.space_after = Pt(6)

            # Book name line
            p2 = doc.add_paragraph()
            r2 = p2.add_run(line.upper())
            r2.bold = True
            r2.font.size = BOOK_NAME_SIZE
            p2.paragraph_format.alignment = 1
            p2.paragraph_format.space_after = Pt(24)
            continue

        # -------- Chapter numbers --------
        if CHAPTER_ONLY.match(line):
            pending_chapter = line
            continue

        # -------- Verses --------
        verse_match = VERSE_LINE.match(line)
        if verse_match:
            verse_num, spacer, verse_text = verse_match.groups()
            p = doc.add_paragraph()

            if verse_num == "1" and pending_chapter:
                r = p.add_run(pending_chapter)
                r.bold = True
                r.font.size = CHAPTER_FONT_SIZE
                p.add_run(spacer + verse_text)
                pending_chapter = None
            else:
                p.add_run(line)
            continue

        # -------- Section headings --------
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(12.5)

# ================= SAVE =================
doc.save(OUTPUT_FILE)
print("Finished: KJV_Cleaned_Final.docx created")
