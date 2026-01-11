import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns

INPUT_FILE = "kjv_formatted.txt"
OUTPUT_FILE = "KJV_Cleaned_Final.docx"

doc = Document()

# ---------------- BASE STYLE ----------------
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

CHAPTER_FONT_SIZE = Pt(20)

# ---------------- REGEX ----------------
ITALICS_PATTERN = re.compile(r'_(.*?)_')
JUNK_START = re.compile(r'^[\u25A0\u25A1\uFFFD\s]+')
REMOVE_KJV_ONLINE = re.compile(r'KJV[\s_]*Online', re.IGNORECASE)

CHAPTER_ONLY = re.compile(r'^\d+$')
VERSE_LINE = re.compile(r'^(\d+)([\u202F\u00A0\s]+)(.*)')
UNDERSCORE_ONLY = re.compile(r'^_+$')
BOOK_TITLE_DASHED = re.compile(r'^-+\s*(.+?)\s*-+$')
BOOK_CHAPTER_LINE = re.compile(r'^(.+?)\s+(\d+)$')  # Added missing regex

# ---------------- TOC ----------------
def add_toc(document):
    p = document.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(ns.qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-1" \\h \\z \\u'
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

# Insert TOC
add_toc(doc)
doc.add_page_break()

pending_chapter = None

with open(INPUT_FILE, 'r', encoding='utf-8') as f:
    for raw in f:
        line = raw.rstrip()

        # Remove KJV Online
        line = REMOVE_KJV_ONLINE.sub('', line)
        if not line.strip():
            continue

        # Remove junk
        line = JUNK_START.sub('', line)

        # Remove underscore spacer lines
        if UNDERSCORE_ONLY.match(line):
            continue

        # Book titles like --- Genesis ---
        book_title = BOOK_TITLE_DASHED.match(line)
        if book_title:
            p = doc.add_paragraph()
            r = p.add_run(book_title.group(1))
            r.bold = True
            r.font.size = Pt(22)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(16)
            continue

        # Chapter number alone (e.g. "16")
        if CHAPTER_ONLY.match(line):
            pending_chapter = line
            continue

        # Book + chapter (e.g. "Genesis 16", "John 3")
        book_chapter = BOOK_CHAPTER_LINE.match(line)
        if book_chapter:
            pending_chapter = book_chapter.group(2)  # only the number
            continue

        # Verse lines
        verse_match = VERSE_LINE.match(line)
        if verse_match:
            verse_num, spacer, verse_text = verse_match.groups()
            p = doc.add_paragraph()

            # Inject chapter number into verse 1
            if verse_num == "1" and pending_chapter:
                chap = p.add_run(pending_chapter)
                chap.bold = True
                chap.font.size = CHAPTER_FONT_SIZE
                p.add_run(spacer + verse_text)
                pending_chapter = None
            else:
                p.add_run(line)
            continue

        # Section headings
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(12.5)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)

doc.save(OUTPUT_FILE)
print("Finished: KJV_Cleaned_Final.docx created")