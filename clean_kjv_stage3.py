import re
from docx import Document
from docx.shared import Pt

INPUT_FILE = "kjv_formatted.txt"
OUTPUT_FILE = "KJV_Cleaned_Final.docx"

doc = Document()

# Base font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# Regex patterns
ITALICS_PATTERN = re.compile(r'_(.*?)_')
JUNK_START = re.compile(r'^[\u25A0\u25A1\uFFFD\s]+')
REMOVE_KJV_ONLINE = re.compile(r'KJV\s*Online', re.IGNORECASE)

CHAPTER_ONLY = re.compile(r'^\d+$')
VERSE_LINE = re.compile(r'^\d+[\u202F\u00A0\s]')  # verse number + space

BOOK_NAMES = re.compile(
    r'^(Genesis|Exodus|Leviticus|Numbers|Deuteronomy|Joshua|Judges|Ruth|'
    r'1 Samuel|2 Samuel|1 Kings|2 Kings|1 Chronicles|2 Chronicles|'
    r'Ezra|Nehemiah|Esther|Job|Psalms|Proverbs|Ecclesiastes|Song of Solomon|'
    r'Isaiah|Jeremiah|Lamentations|Ezekiel|Daniel|Hosea|Joel|Amos|Obadiah|'
    r'Jonah|Micah|Nahum|Habakkuk|Zephaniah|Haggai|Zechariah|Malachi|'
    r'Matthew|Mark|Luke|John|Acts|Romans|1 Corinthians|2 Corinthians|'
    r'Galatians|Ephesians|Philippians|Colossians|'
    r'1 Thessalonians|2 Thessalonians|1 Timothy|2 Timothy|Titus|Philemon|'
    r'Hebrews|James|1 Peter|2 Peter|1 John|2 John|3 John|Jude|Revelation)$'
)

with open(INPUT_FILE, 'r', encoding='utf-8') as f:
    for raw_line in f:
        line = raw_line.rstrip()

        # Remove "KJV Online"
        line = REMOVE_KJV_ONLINE.sub('', line)
        if not line.strip():
            continue

        # Remove junk squares
        line = JUNK_START.sub('', line)

        # Chapter numbers (big + bold)
        if CHAPTER_ONLY.match(line):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(20)
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            continue

        # Book names (bold, larger)
        if BOOK_NAMES.match(line):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(16)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(10)
            continue

        # SECTION HEADINGS (THIS IS THE KEY CHANGE)
        if not VERSE_LINE.match(line):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(12.5)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            continue

        # Verse lines with italics
        p = doc.add_paragraph()
        last = 0
        for m in ITALICS_PATTERN.finditer(line):
            p.add_run(line[last:m.start()])
            ital = p.add_run(m.group(1))
            ital.italic = True
            last = m.end()
        p.add_run(line[last:])

doc.save(OUTPUT_FILE)

print("Finished: KJV_Cleaned_Final.docx created")
