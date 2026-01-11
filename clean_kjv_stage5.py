import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns

INPUT_FILE = "kjv_formatted.txt"
OUTPUT_FILE = "KJV_Cleaned_Final.docx"

doc = Document()

# ---------------- BASE STYLE ----------------
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# ---------------- REGEX ----------------
ITALICS_PATTERN = re.compile(r'_(.*?)_')
JUNK_START = re.compile(r'^[\u25A0\u25A1\uFFFD\s]+')
REMOVE_KJV_ONLINE = re.compile(r'KJV[\s_]*Online', re.IGNORECASE)

CHAPTER_ONLY = re.compile(r'^\d+$')
BOOK_CHAPTER = re.compile(
    r'^(Genesis|Exodus|Leviticus|Numbers|Deuteronomy|Joshua|Judges|Ruth|'
    r'1 Samuel|2 Samuel|1 Kings|2 Kings|1 Chronicles|2 Chronicles|'
    r'Ezra|Nehemiah|Esther|Job|Psalms|Proverbs|Ecclesiastes|Song of Solomon|'
    r'Isaiah|Jeremiah|Lamentations|Ezekiel|Daniel|Hosea|Joel|Amos|Obadiah|'
    r'Jonah|Micah|Nahum|Habakkuk|Zephaniah|Haggai|Zechariah|Malachi|'
    r'Matthew|Mark|Luke|John|Acts|Romans|1 Corinthians|2 Corinthians|'
    r'Galatians|Ephesians|Philippians|Colossians|'
    r'1 Thessalonians|2 Thessalonians|1 Timothy|2 Timothy|Titus|Philemon|'
    r'Hebrews|James|1 Peter|2 Peter|1 John|2 John|3 John|Jude|Revelation)'
    r'\s+(\d+)$'
)

BOOK_TITLE_DASHED = re.compile(r'^-+\s*(.+?)\s*-+$')
VERSE_LINE = re.compile(r'^\d+[\u202F\u00A0\s]')
UNDERSCORE_ONLY = re.compile(r'^_+$')

# ---------------- TOC ----------------
def add_toc(document):
    p = document.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(ns.qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-1" \\h \\z \\u'
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

# Insert TOC at top
add_toc(doc)
doc.add_page_break()

# ---------------- PROCESS TEXT ----------------
with open(INPUT_FILE, 'r', encoding='utf-8') as f:
    for raw in f:
        line = raw.rstrip()

        # Remove KJV Online
        line = REMOVE_KJV_ONLINE.sub('', line)
        if not line.strip():
            continue

        # Remove junk symbols
        line = JUNK_START.sub('', line)

        # Remove stray underscore spacer lines
        if UNDERSCORE_ONLY.match(line):
            continue

        # Book titles like --- Genesis ---
        book_title = BOOK_TITLE_DASHED.match(line)
        if book_title:
            p = doc.add_paragraph()
            r = p.add_run(book_title.group(1))
            r.bold = True
            r.font.size = Pt(22)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(16)
            continue

        # Book + chapter → chapter only
        bc = BOOK_CHAPTER.match(line)
        if bc:
            p = doc.add_paragraph()
            r = p.add_run(bc.group(2))
            r.bold = True
            r.font.size = Pt(20)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            continue

        # Chapter alone
        if CHAPTER_ONLY.match(line):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(20)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            continue

        # Section headings
        if not VERSE_LINE.match(line):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(12.5)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            continue

        # Verse lines (italics preserved)
        p = doc.add_paragraph()
        last = 0
        for m in ITALICS_PATTERN.finditer(line):
            p.add_run(line[last:m.start()])
            ital = p.add_run(m.group(1))
            ital.italic = True
            last = m.end()
        p.add_run(line[last:])

doc.save(OUTPUT_FILE)
print("Finished: KJV_Cleaned_Final.docx created")
